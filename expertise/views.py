# expertise/views.py

from django.shortcuts import get_object_or_404, render, redirect
from django.urls import reverse_lazy, reverse
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.db.models import Q, Sum
from django.db import transaction
from django.http import HttpResponse
from django.utils import timezone
from datetime import datetime, timedelta
from collections import defaultdict
from decimal import Decimal, ROUND_HALF_UP
from num2words import num2words
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import barcode
import io
import base64
import zipfile
import pandas as pd
from barcode.writer import ImageWriter
from django.views.decorators.http import require_POST
from .models import (
    FicheEvenement,
    PersonnelNavigant,
    CompagnieAerienne,
    Bordereau,
    FactureMedecin,
    Medecin,
    MedecinInvoice,
    MedecinInvoiceLine,
)
from .forms import BordereauSelectionForm
from django.db import models
from django.template.loader import render_to_string
from weasyprint import HTML
from django.contrib.auth.decorators import login_required
from django.utils.decorators import method_decorator
from django.contrib.auth.mixins import LoginRequiredMixin

from django.contrib.auth.views import LoginView, LogoutView
from django.shortcuts import render
from django.contrib.auth.decorators import login_required

@login_required(login_url='/login/')
def accueil(request):
    return render(request, 'expertise/accueil.html')

from django.contrib.auth import logout
from django.shortcuts import redirect

def logout_view(request):
    logout(request)
    return redirect('/login/')

@login_required(login_url='/login/')
def assign_bordereau(request, mois, annee, iata):
    compagnie = get_object_or_404(CompagnieAerienne, iata=iata)
    evenements = FicheEvenement.objects.filter(
        date_evenement__year=annee, date_evenement__month=mois, personnel__compagnie=compagnie
    )

    # Création du bordereau
    date_creation = datetime.today()
    no_bordereau = f"EB{date_creation.day:02d}{mois:02d}{str(annee)[-2:]}{iata}"
    bordereau = Bordereau.objects.create(
        date_bordereau=date_creation,
        no_bordereau=no_bordereau
    )

    # Lier le bordereau aux événements
    for evenement in evenements:
        evenement.no_bordereau = no_bordereau
        evenement.bordereau = bordereau
        evenement.save()

    return redirect('bordereau_view', mois=mois, annee=annee, iata=iata)

# ----- VUES POUR LES PERSONNELS -----
# Pour une classe
class PersonnelListView(LoginRequiredMixin, ListView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_list.html'
    context_object_name = 'personnels'

    def get_queryset(self):
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(nom__icontains=query) |
                Q(prenom__icontains=query) |
                Q(dn__icontains=query)
            )
        return queryset


class PersonnelDetailView(LoginRequiredMixin, DetailView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_detail.html'
    context_object_name = 'personnel'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['evenements'] = self.object.evenements.all()
        return context


# ----- VUES POUR LES EVENEMENTS -----
from django.utils.timezone import now

class FicheEvenementCreateView(LoginRequiredMixin, CreateView):
    model = FicheEvenement
    fields = [
        'date_evenement',
        'cs_cempn', 'date_cempn', 'honoraire_cempn', 'medecin_cempn',
        'cs_oph', 'date_cs_oph', 'honoraire_cs_oph', 'medecin_oph',
        'cs_orl', 'date_cs_orl', 'honoraire_cs_orl', 'medecin_orl',
        'cs_labo', 'date_cs_labo', 'honoraire_cs_labo', 'medecin_labo',
        'cs_lbx', 'date_cs_lbx', 'honoraire_cs_lbx',
        'cs_toxique', 'date_cs_toxique', 'honoraire_cs_toxique',
        'cs_radio', 'date_cs_radio', 'honoraire_cs_radio', 'medecin_radio',
        'frais_dossier', 'quote_part_patient',
        'paiement', 'date_paiement', 'modalite_paiement',
    ]
    template_name = 'expertise/evenement_form.html'

    def get_initial(self):
        date = now().date()  # ou une autre logique si tu veux personnaliser
        return {
            'date_evenement': date,
            'date_cempn': date,
            'date_cs_oph': date,
            'date_cs_orl': date,
            'date_cs_labo': date,
            'date_cs_lbx': date,
            'date_cs_toxique': date,
            'date_cs_radio': date,
            'date_paiement': date
        }

    def form_valid(self, form):
        dn = self.kwargs['dn']
        personnel = get_object_or_404(PersonnelNavigant, dn=dn)
        form.instance.personnel = personnel
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['dn'] = self.kwargs['dn']
        return context

    def get_success_url(self):
        return reverse_lazy('personnel_detail', kwargs={'dn': self.object.personnel.dn})




class FactureView(LoginRequiredMixin, DetailView):
    model = FicheEvenement
    template_name = 'expertise/facture.html'
    context_object_name = 'evenement'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        invoice_number = self.object.no_facture
        print(f"DEBUG: Numéro de facture = {invoice_number}")

        if invoice_number:
            try:
                Code128 = barcode.get_barcode_class('code128')
                barcode_instance = Code128(invoice_number, writer=ImageWriter())
                buffer = io.BytesIO()
                barcode_instance.write(buffer)
                barcode_base64 = base64.b64encode(buffer.getvalue()).decode('utf-8')
                context['barcode'] = barcode_base64
            except Exception as e:
                print(f"Erreur lors de la génération du code-barres : {e}")

        return context


# ----- VUE DU BORDEREAU -----
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words

from .models import CompagnieAerienne, FicheEvenement, Bordereau

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from num2words import num2words

from .models import CompagnieAerienne, FicheEvenement, Bordereau
from docx.shared import RGBColor

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.db import transaction
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from num2words import num2words

from .models import CompagnieAerienne, FicheEvenement, Bordereau

@login_required(login_url='/login/')
def download_bordereau(request, mois, annee, iata):
    compagnie = get_object_or_404(CompagnieAerienne, iata=iata)
    evenements = FicheEvenement.objects.filter(
        date_evenement__year=annee,
        date_evenement__month=mois,
        personnel__compagnie=compagnie
    )

    date_creation = datetime.today()
    no_bordereau = f"EB{date_creation.day:02d}{mois:02d}{str(annee)[-2:]}{iata}"

    bordereau, created = Bordereau.objects.get_or_create(
        no_bordereau=no_bordereau,
        defaults={"date_bordereau": date_creation}
    )

    with transaction.atomic():
        evenements.update(bordereau=bordereau)

    doc = Document()

    para = doc.add_heading('Centre Médical du Personnel Navigant de Polynésie française', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Dr. Christian Hellec', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('BP 380697 - 98718 Punaauia - Tahiti', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Polynésie Française', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('mel : cmpnpf@gmail.com | Tel : +689.87.77.05.18 | Tel : +689.87.71.50.90', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Bordereau de dépôt de factures', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('--------------------', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Date de création : {date_creation.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Numéro du bordereau : {no_bordereau}")
    doc.add_paragraph(f"Compagnie aérienne : {compagnie.nom} ({compagnie.iata})")

    # Tableau principal
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdrs = ["Numéro de facture", "DN", "Nom", "Prénom", "Total (XPF)", "Paiement"]
    for i, header in enumerate(hdrs):
        cell = table.rows[0].cells[i]
        cell.text = header
        para = cell.paragraphs[0]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.runs[0]
        run.bold = True

    total_general = 0

    for e in evenements:
        total_general += e.total or 0
        row = table.add_row().cells
        data = [
            e.no_facture or "N/A",
            e.personnel.dn,
            e.personnel.nom,
            e.personnel.prenom,
            f"{e.total or 0:,} XPF",
            "Payé" if e.paiement else "Non payé"
        ]
        for i, val in enumerate(data):
            para = row[i].paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.add_run(val)

    total_lettres = num2words(total_general, lang='fr').capitalize()
    doc.add_paragraph(
        f"\nNombre de factures : {evenements.count()} | Total général : {total_general:,} XPF ({total_lettres})"
    )
    doc.add_paragraph("Dr. Christian HELLEC")
    doc.add_paragraph("IBAN : FR76 1223 9000 0162 2887 0100 014")

    # Factures individuelles
    for e in evenements:
        doc.add_page_break()
        para = doc.add_heading('Centre Médical du Personnel Navigant de Polynésie française', level=1)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('Dr. Christian Hellec', level=1)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('BP 380697 - 98718 Punaauia - Tahiti', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('Polynésie Française', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('mel : cmpnpf@gmail.com | Tel : +689.87.77.05.18 | Tel : +689.87.71.50.90', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('Facture Individuelle', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading(f"{compagnie.nom}", level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para = doc.add_heading('--------------------', level=2)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")
        doc.add_paragraph(f"Date : {e.date_evenement.strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Numéro de facture : {e.no_facture or 'N/A'}")

        doc.add_heading("Informations du patient", level=2)
        doc.add_paragraph(f"Nom : {e.personnel.nom}")
        doc.add_paragraph(f"Prénom : {e.personnel.prenom}")
        doc.add_paragraph(f"DN : {e.personnel.dn}")
        if e.personnel.date_de_naissance:
            doc.add_paragraph(f"Date de naissance : {e.personnel.date_de_naissance.strftime('%d/%m/%Y')}")

        doc.add_heading("Détails des actes", level=2)
        actes = [
            ("CEMPN/Pf", e.cs_cempn, e.date_cempn, e.medecin_cempn, e.honoraire_cempn),
            ("Ophtalmologie", e.cs_oph, e.date_cs_oph, e.medecin_oph, e.honoraire_cs_oph),
            ("ORL", e.cs_orl, e.date_cs_orl, e.medecin_orl, e.honoraire_cs_orl),
            ("Biologie sanguine", e.cs_labo, e.date_cs_labo, e.medecin_labo, e.honoraire_cs_labo),
            ("Biologie urinaire", e.cs_lbx, e.date_cs_lbx, e.medecin_labo, e.honoraire_cs_lbx),
            ("Toxicologie", e.cs_toxique, e.date_evenement, e.medecin_labo, e.honoraire_cs_toxique),
            ("Radiologie", e.cs_radio, e.date_cs_radio, e.medecin_radio, e.honoraire_cs_radio),
            ("Frais de dossier", e.cs_cempn, e.date_cempn, e.medecin_cempn, e.frais_dossier),
        ]

        # Optionnel si ce champ existe dans ton modèle :
        #if hasattr(e, 'honoraire_cs_toxique'):
            #actes.append(("TOXICO", True, e.date_evenement, e.medecin_labo, e.honoraire_cs_toxique))

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        headers = ["Date", "Acte", "Médecin", "Montant (XPF)", "Quote-part"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            para.runs[0].bold = True

        for libelle, actif, date_acte, medecin, montant in actes:
            if actif:
                row = table.add_row().cells
                data = [
                    date_acte.strftime('%d/%m/%Y') if date_acte else "-",
                    libelle,
                    f"{medecin.prenom} {medecin.nom}" if medecin else "-",
                    f"{montant or 0:,} XPF",
                    f"{e.paye_par_patient or 0:,} XPF" if e.quote_part_patient else "-"
                ]
                for i, val in enumerate(data):
                    para = row[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    para.add_run(val)

        doc.add_paragraph(f"\n💰 Total : {e.total or 0:,} XPF")
        doc.add_paragraph(f"🧾 Payé par le patient : {e.paye_par_patient or 0:,} XPF")
        doc.add_paragraph("Dr. Christian HELLEC")
        doc.add_paragraph("IBAN : FR76 1223 9000 0162 2887 0100 014")

    # Générer le fichier en réponse HTTP
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Bordereau_{no_bordereau}.docx"'
    doc.save(response)
    return response



# ----- LISTE DES BORDEREAUX -----
from django.shortcuts import render
from .models import Bordereau
from django.db.models import Sum
from .models import Bordereau

@login_required(login_url='/login/')
def liste_bordereaux(request):
    bordereaux = Bordereau.objects.all().prefetch_related('evenements')

    for bordereau in bordereaux:
        bordereau.total_general = (
            bordereau.evenements.aggregate(Sum('total'))['total__sum'] or 0
        )
        bordereau.nb_factures = bordereau.evenements.count()

    return render(request, 'expertise/liste_bordereaux.html', {
        'bordereaux': bordereaux
    })





# ----- SELECTION DE BORDEREAU -----
def bordereau_selection_view(request):
    if request.method == 'POST':
        form = BordereauSelectionForm(request.POST)
        if form.is_valid():
            mois = form.cleaned_data['mois']
            annee = form.cleaned_data['annee']
            compagnie = form.cleaned_data['compagnie']
            iata = compagnie.iata
            return redirect('bordereau_detail', annee=annee, mois=mois, iata=iata)
    else:
        form = BordereauSelectionForm()

    return render(request, 'expertise/selection_bordereau.html', {'form': form})

# expertise/views.py

from django.shortcuts import render, get_object_or_404, redirect
from django.urls import reverse_lazy, reverse
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.db.models import Q
from .models import PersonnelNavigant, FicheEvenement, CompagnieAerienne
from .forms import BordereauSelectionForm
from .utils import nombre_en_lettres
from num2words import num2words
from datetime import datetime
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.db import transaction
import io
import base64
import barcode
from barcode.writer import ImageWriter

# ----- VUES POUR LES PERSONNELS -----

class PersonnelListView(LoginRequiredMixin, ListView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_list.html'
    context_object_name = 'personnels'

    def get_queryset(self):
        queryset = super().get_queryset()
        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(nom__icontains=query) |
                Q(prenom__icontains=query) |
                Q(dn__icontains=query)
            )
        return queryset


class PersonnelDetailView(LoginRequiredMixin, DetailView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_detail.html'
    context_object_name = 'personnel'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['evenements'] = self.object.evenements.all()
        return context


class PersonnelCreateView(LoginRequiredMixin, CreateView):
    model = PersonnelNavigant
    fields = ['dn', 'nom', 'prenom', 'date_de_naissance', 'sexe', 'statut_pn', 'compagnie']
    template_name = 'expertise/personnel_form.html'
    success_url = reverse_lazy('personnel_list')


class PersonnelUpdateView(UpdateView):
    model = PersonnelNavigant
    fields = ['dn', 'nom', 'prenom', 'date_de_naissance', 'sexe', 'statut_pn', 'compagnie']
    template_name = 'expertise/personnel_form.html'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'
    success_url = reverse_lazy('personnel_list')


class PersonnelDeleteView(DeleteView):
    model = PersonnelNavigant
    template_name = 'expertise/personnel_confirm_delete.html'
    slug_field = 'dn'
    slug_url_kwarg = 'dn'
    success_url = reverse_lazy('personnel_list')

from django.views.generic import UpdateView
from .models import FicheEvenement
from django.urls import reverse_lazy

class FicheEvenementUpdateView(UpdateView):
    model = FicheEvenement
    fields = [
        'date_evenement',
        'cs_cempn', 'date_cempn', 'honoraire_cempn',
        'cs_oph', 'date_cs_oph', 'honoraire_cs_oph',
        'cs_orl', 'date_cs_orl', 'honoraire_cs_orl',
        'cs_labo', 'date_cs_labo', 'honoraire_cs_labo',
        'cs_lbx', 'date_cs_lbx', 'honoraire_cs_lbx',
        'cs_toxique', 'date_cs_toxique', 'honoraire_cs_toxique',
        'cs_radio', 'date_cs_radio', 'honoraire_cs_radio',
        'medecin_cempn', 'medecin_oph',
        'medecin_orl', 'medecin_radio', 'medecin_labo',
        'frais_dossier', 'quote_part_patient',
        'paiement', 'date_paiement', 'modalite_paiement',
    ]
    template_name = 'expertise/evenement_form.html'

    def get_success_url(self):
        if self.object.personnel and self.object.personnel.dn:
            return reverse_lazy('personnel_detail', kwargs={'dn': self.object.personnel.dn})
        else:
            return reverse_lazy('personnel_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.object.personnel:
            context['dn'] = self.object.personnel.dn
        return context


from django.views.generic import DeleteView
from django.urls import reverse_lazy
from .models import FicheEvenement

from django.urls import reverse_lazy
from django.views.generic import DeleteView
from .models import FicheEvenement

class FicheEvenementDeleteView(DeleteView):
    model = FicheEvenement
    template_name = 'expertise/evenement_confirm_delete.html'

    def get_success_url(self):
        # if the deleted event had a linked personnel, go back to their detail page
        if self.object.personnel:
            return reverse_lazy('personnel_detail', kwargs={'dn': self.object.personnel.dn})
        # otherwise, fall back to a generic list
        return reverse_lazy('personnel_list')



from django.shortcuts import render, get_object_or_404
from datetime import datetime
from .models import FicheEvenement, CompagnieAerienne
from .utils import nombre_en_lettres  # Assure-toi que ce module existe

def bordereau_view(request, annee, mois, iata):
    compagnie = get_object_or_404(CompagnieAerienne, iata=iata)

    evenements = FicheEvenement.objects.filter(
        date_evenement__year=annee,
        date_evenement__month=mois,
        personnel__compagnie=compagnie
    )

    date_bordereau = datetime.today().strftime('%d/%m/%Y')
    no_bordereau = f"EB{datetime.today().day:02d}{mois:02d}{str(annee)[-2:]}{iata}"
    total_global = sum(e.total for e in evenements)
    total_global_lettres = nombre_en_lettres(total_global)

    # Création ou récupération du bordereau
    bordereau, _ = Bordereau.objects.get_or_create(
    no_bordereau=no_bordereau,
    defaults={'date_bordereau': datetime.today()}
)

    # Mise à jour de chaque événement
    with transaction.atomic():
        for e in evenements:
            e.bordereau = bordereau
            e.save()


    return render(request, "expertise/bordereau.html", {
        "evenements": evenements,
        "mois": mois,
        "annee": annee,
        "iata": iata,
        "compagnie": compagnie,
        "date_bordereau": date_bordereau,
        "no_bordereau": no_bordereau,
        "nombre_factures": evenements.count(),
        "total_global": total_global,
        "total_global_lettres": total_global_lettres,
    })
# views.py

from .models import FicheEvenement

def bordereau_factures(request, no_bordereau):
    evenements = FicheEvenement.objects.filter(bordereau__no_bordereau=no_bordereau)
    return render(request, 'expertise/factures_bordereau.html', {
        'evenements': evenements,
        'no_bordereau': no_bordereau
    })

from django.shortcuts import redirect, get_object_or_404
from .models import Bordereau
@login_required(login_url='/login/')
@require_POST
def toggle_virement(request, id):
    bordereau = get_object_or_404(Bordereau, id=id)
    bordereau.virement = not bordereau.virement
    bordereau.save(update_fields=['virement'])

    evenements_qs = bordereau.evenements.select_related(
        'medecin_cempn', 'medecin_oph', 'medecin_orl',
        'medecin_radio', 'medecin_labo', 'personnel__compagnie'
    )

    evenements = list(evenements_qs)

    if evenements:
        if bordereau.virement:
            today = timezone.now().date()
            for evenement in evenements:
                updated_fields = []
                if not evenement.paiement:
                    evenement.paiement = True
                    updated_fields.append('paiement')
                if not evenement.date_paiement:
                    evenement.date_paiement = today
                    updated_fields.append('date_paiement')
                if updated_fields:
                    evenement.save(update_fields=updated_fields)
        else:
            for evenement in evenements:
                updated_fields = []
                if evenement.paiement:
                    evenement.paiement = False
                    updated_fields.append('paiement')
                if evenement.date_paiement is not None:
                    evenement.date_paiement = None
                    updated_fields.append('date_paiement')
                if updated_fields:
                    evenement.save(update_fields=updated_fields)

    FactureMedecin.objects.filter(bordereau=bordereau).delete()

    if bordereau.virement and evenements:
        honoraires_medecins = {}

        for evenement in evenements:
            for champ, montant in [
                (evenement.medecin_cempn, evenement.honoraire_cempn),
                (evenement.medecin_oph, evenement.honoraire_cs_oph),
                (evenement.medecin_orl, evenement.honoraire_cs_orl),
                (evenement.medecin_radio, evenement.honoraire_cs_radio),
                (evenement.medecin_labo, evenement.honoraire_cs_labo),
                (evenement.medecin_labo, evenement.honoraire_cs_lbx),
                (evenement.medecin_labo, evenement.honoraire_cs_toxique),
            ]:
                if champ:
                    honoraires_medecins.setdefault(champ, 0)
                    honoraires_medecins[champ] += montant or 0

        factures_medecins = [
            FactureMedecin(medecin=medecin, bordereau=bordereau, montant=total)
            for medecin, total in honoraires_medecins.items()
        ]

        if factures_medecins:
            FactureMedecin.objects.bulk_create(factures_medecins)

    next_url = request.POST.get('next')
    if next_url:
        return redirect(next_url)
    return redirect('liste_bordereaux')



def factures_medecins_bordereau(request, no_bordereau):
    bordereau = get_object_or_404(Bordereau, no_bordereau=no_bordereau)
    factures = FactureMedecin.objects.filter(bordereau=bordereau)

    for facture in factures:
        if facture.medecin.nom.lower() == "hellec":
            facture.redevance = Decimal('0')
        else:
            facture.redevance = round(facture.montant * Decimal('0.06'), 0)
        facture.montant_net = round(facture.montant - facture.redevance, 0)

    return render(request, 'expertise/factures_medecins_bordereau.html', {
        'bordereau': bordereau,
        'factures': factures,
    })


@login_required(login_url='/login/')
def export_evenements_excel(request):
    evenements = (
        FicheEvenement.objects
        .select_related('personnel__compagnie')
        .order_by('date_evenement', 'personnel__nom', 'personnel__prenom')
    )

    rows = []
    for evenement in evenements:
        personnel = evenement.personnel
        compagnie = personnel.compagnie if personnel else None

        rows.append({
            'Date': evenement.date_evenement.strftime('%d/%m/%Y') if evenement.date_evenement else '',
            'Nom': personnel.nom if personnel else '',
            'Prenom': personnel.prenom if personnel else '',
            'Date de naissance': personnel.date_de_naissance.strftime('%d/%m/%Y') if personnel and personnel.date_de_naissance else '',
            'Numero de facture': evenement.no_facture or '',
            'Total de la facture': evenement.total or 0,
            'Statut facture': 'Payée' if evenement.paiement else 'Non payée',
            'Statut': personnel.get_statut_pn_display() if personnel and personnel.statut_pn else '',
            'Compagnie': compagnie.nom if compagnie else '',
        })

    df = pd.DataFrame(rows, columns=[
        'Date',
        'Nom',
        'Prenom',
        'Date de naissance',
        'Numero de facture',
        'Total de la facture',
        'Statut facture',
        'Statut',
        'Compagnie',
    ])

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Factures')

    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="export_factures.xlsx"'
    return response


@login_required(login_url='/login/')
def facture_search(request):
    search_query = (request.GET.get('q') or '').strip()
    selected_number = (request.GET.get('facture') or '').strip()

    matches_qs = (
        FicheEvenement.objects
        .select_related('personnel__compagnie')
        .order_by('no_facture')
    )

    if search_query:
        matches_qs = matches_qs.filter(no_facture__icontains=search_query)

    matches = list(matches_qs)

    selected_facture = None
    if selected_number:
        selected_facture = next((f for f in matches if f.no_facture == selected_number), None)
    elif search_query and len(matches) == 1:
        selected_facture = matches[0]

    statut_label = None
    if selected_facture:
        statut_label = 'Payée' if selected_facture.paiement else 'Non payée'

    return render(request, 'expertise/facture_search.html', {
        'search_query': search_query,
        'matches': matches,
        'selected_facture': selected_facture,
        'statut_label': statut_label,
    })


@login_required(login_url='/login/')
@require_POST
def toggle_facture_paiement(request, pk):
    facture = get_object_or_404(FicheEvenement, pk=pk)
    facture.paiement = not facture.paiement
    if facture.paiement and not facture.date_paiement:
        facture.date_paiement = timezone.now().date()
    elif not facture.paiement:
        facture.date_paiement = None

    facture.save(update_fields=['paiement', 'date_paiement'])

    next_url = request.POST.get('next') or reverse('facture_search')
    return redirect(next_url)


@login_required(login_url='/login/')
def facture_par_compagnie(request):
    compagnies = CompagnieAerienne.objects.order_by('nom')
    selected_compagnie_id = (request.GET.get('compagnie') or '').strip()
    selected_facture_no = (request.GET.get('facture') or '').strip()

    selected_compagnie = None
    if selected_compagnie_id:
        selected_compagnie = compagnies.filter(pk=selected_compagnie_id).first()

    factures_qs = FicheEvenement.objects.none()
    if selected_compagnie:
        factures_qs = (
            FicheEvenement.objects
            .select_related('personnel__compagnie')
            .filter(personnel__compagnie=selected_compagnie)
            .order_by('no_facture')
        )

    factures = list(factures_qs)

    selected_facture = None
    if selected_facture_no:
        selected_facture = next((f for f in factures if f.no_facture == selected_facture_no), None)
    elif factures and len(factures) == 1:
        selected_facture = factures[0]

    statut_label = None
    if selected_facture:
        statut_label = 'Payée' if selected_facture.paiement else 'Non payée'

    return render(request, 'expertise/facture_par_compagnie.html', {
        'compagnies': compagnies,
        'selected_compagnie': selected_compagnie,
        'factures': factures,
        'selected_facture': selected_facture,
        'statut_label': statut_label,
    })


@login_required(login_url='/login/')
def bordereau_par_compagnie(request):
    compagnies = CompagnieAerienne.objects.order_by('nom')
    selected_compagnie_id = (request.GET.get('compagnie') or '').strip()
    selected_bordereau_id = (request.GET.get('bordereau') or '').strip()

    selected_compagnie = None
    if selected_compagnie_id:
        selected_compagnie = compagnies.filter(pk=selected_compagnie_id).first()

    bordereaux_qs = Bordereau.objects.none()
    if selected_compagnie:
        bordereaux_qs = Bordereau.objects.filter(
            evenements__personnel__compagnie=selected_compagnie
        ).distinct().order_by('-date_bordereau', 'no_bordereau')

    bordereaux = list(bordereaux_qs)

    selected_bordereau = None
    if selected_bordereau_id:
        selected_bordereau = next((b for b in bordereaux if str(b.id) == selected_bordereau_id), None)
    elif len(bordereaux) == 1:
        selected_bordereau = bordereaux[0]

    factures = []
    bordereau_total = 0
    if selected_bordereau:
        factures_qs = (
            FicheEvenement.objects
            .select_related('personnel__compagnie')
            .filter(bordereau=selected_bordereau)
            .order_by('no_facture')
        )
        bordereau_total = factures_qs.aggregate(total_sum=Sum('total'))['total_sum'] or 0
        factures = list(factures_qs)

    return render(request, 'expertise/bordereau_par_compagnie.html', {
        'compagnies': compagnies,
        'selected_compagnie': selected_compagnie,
        'bordereaux': bordereaux,
        'selected_bordereau': selected_bordereau,
        'factures': factures,
        'bordereau_total': bordereau_total,
    })


ACTE_CONFIGS = [
    ('medecin_cempn', 'Consultation CEMPN', 'cs_cempn', 'date_cempn', 'honoraire_cempn', 'CEMPN'),
    ('medecin_oph', 'Consultation OPH', 'cs_oph', 'date_cs_oph', 'honoraire_cs_oph', 'OPH'),
    ('medecin_orl', 'Consultation ORL', 'cs_orl', 'date_cs_orl', 'honoraire_cs_orl', 'ORL'),
    ('medecin_radio', 'Consultation Radio', 'cs_radio', 'date_cs_radio', 'honoraire_cs_radio', 'RADIO'),
    ('medecin_labo', 'Biologie sanguine', 'cs_labo', 'date_cs_labo', 'honoraire_cs_labo', 'LABO'),
    ('medecin_labo', 'Biologie urinaire', 'cs_lbx', 'date_cs_lbx', 'honoraire_cs_lbx', 'LABOX'),
    ('medecin_labo', 'Recherche toxique', 'cs_toxique', 'date_cs_toxique', 'honoraire_cs_toxique', 'TOX'),
]


def _get_redevance_rate(medecin):
    if not medecin:
        return Decimal('0.06')

    nom = (medecin.nom or '').strip().upper()
    if nom == 'HELLEC':
        return Decimal('0.00')

    specialite = (medecin.specialite or '').lower()
    if 'radiolog' in specialite:
        return Decimal('0.00')
    if 'laboratoire' in specialite or 'labo' in specialite:
        return Decimal('0.10')

    return Decimal('0.06')


def _generate_medecin_invoice_number(medecin, emission_date):
    initiale = (medecin.nom.strip()[0] if medecin.nom else medecin.prenom[:1]).upper()
    base = f"{emission_date.year}-{initiale}-{emission_date.strftime('%d%m%Y')}"
    candidate = base
    suffix = 2
    while MedecinInvoice.objects.filter(number=candidate).exists():
        candidate = f"{base}-{suffix}"
        suffix += 1
    return candidate


def _create_medecin_invoice(medecin, entries):
    emission_date = timezone.now().date()
    number = _generate_medecin_invoice_number(medecin, emission_date)

    total_brut = sum(Decimal(entry.get('montant_brut', Decimal('0'))) for entry in entries)
    total_redevance = sum(Decimal(entry.get('redevance', Decimal('0'))) for entry in entries)
    total_net = sum(Decimal(entry.get('montant_net', Decimal('0'))) for entry in entries)

    invoice = MedecinInvoice.objects.create(
        medecin=medecin,
        number=number,
        emission_date=emission_date,
        total_brut=total_brut.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP),
        total_redevance=total_redevance.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP),
        total_net=total_net.quantize(Decimal('0.01'), rounding=ROUND_HALF_UP),
    )

    for entry in entries:
        date_acte = entry.get('date') or emission_date
        MedecinInvoiceLine.objects.create(
            invoice=invoice,
            evenement_id=entry['event_id'],
            act_code=entry['act_code'],
            act_label=entry['acte'],
            date_acte=date_acte,
            patient_nom=entry.get('patient_nom', ''),
            patient_prenom=entry.get('patient_prenom', ''),
            facture_no=entry.get('facture_no'),
            montant_brut=Decimal(entry.get('montant_brut', Decimal('0'))).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP),
            redevance=Decimal(entry.get('redevance', Decimal('0'))).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP),
            montant_net=Decimal(entry.get('montant_net', Decimal('0'))).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP),
        )

    return invoice


def _collect_medecin_histories(target_medecin=None):
    history_map = defaultdict(list)

    medecins_qs = Medecin.objects.order_by('nom', 'prenom')
    if target_medecin:
        medecins_qs = medecins_qs.filter(pk=target_medecin.pk)

    medecins = list(medecins_qs)
    if not medecins:
        return history_map

    medecin_ids = [med.id for med in medecins]

    evenements_qs = (
        FicheEvenement.objects
        .select_related(
            'personnel', 'bordereau',
            'medecin_cempn', 'medecin_oph', 'medecin_orl',
            'medecin_radio', 'medecin_labo'
        )
        .filter(
            Q(medecin_cempn_id__in=medecin_ids)
            | Q(medecin_oph_id__in=medecin_ids)
            | Q(medecin_orl_id__in=medecin_ids)
            | Q(medecin_radio_id__in=medecin_ids)
            | Q(medecin_labo_id__in=medecin_ids)
        )
    )

    evenements = list(evenements_qs)
    if not evenements:
        return history_map

    event_ids = [e.id for e in evenements]

    lines_map = {}
    if event_ids:
        lignes = (
            MedecinInvoiceLine.objects
            .select_related('invoice')
            .filter(
                evenement_id__in=event_ids,
                invoice__medecin_id__in=medecin_ids
            )
        )
        lines_map = {
            (ligne.evenement_id, ligne.act_code): ligne
            for ligne in lignes
        }

    for evenement in evenements:
        for field_name, act_label, flag_field, date_field, amount_field, act_code in ACTE_CONFIGS:
            medecin = getattr(evenement, field_name)
            if not medecin or medecin.id not in medecin_ids:
                continue
            if not getattr(evenement, flag_field, False):
                continue

            act_date = getattr(evenement, date_field, None) or evenement.date_evenement
            montant_brut = Decimal(getattr(evenement, amount_field, 0) or 0)

            ligne = lines_map.get((evenement.id, act_code))

            if ligne:
                redevance = ligne.redevance
                montant_net = ligne.montant_net
                invoice_number = ligne.invoice.number
                invoice_id = ligne.invoice_id
                invoice_date = ligne.invoice.emission_date
            else:
                rate = _get_redevance_rate(medecin)
                redevance = (montant_brut * rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                montant_net = (montant_brut - redevance).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
                invoice_number = ''
                invoice_id = None
                invoice_date = None

            history_map[medecin.id].append({
                'acte': act_label,
                'act_code': act_code,
                'date': act_date,
                'event_id': evenement.id,
                'facture_no': evenement.no_facture,
                'facture_id': evenement.pk,
                'paiement': evenement.paiement,
                'patient_nom': evenement.personnel.nom if evenement.personnel else '',
                'patient_prenom': evenement.personnel.prenom if evenement.personnel else '',
                'montant_brut': montant_brut,
                'montant_net': montant_net,
                'redevance': redevance,
                'invoice_number': invoice_number,
                'invoice_id': invoice_id,
                'invoice_date': invoice_date,
            })

    return history_map


@login_required(login_url='/login/')
def intervenants_list(request):
    medecins = list(Medecin.objects.order_by('nom', 'prenom'))
    history_map = _collect_medecin_histories()

    medecin_rows = []
    for medecin in medecins:
        history = history_map.get(medecin.id, [])
        history.sort(key=lambda item: item['date'] or datetime.min, reverse=True)
        medecin_rows.append({
            'medecin': medecin,
            'history': history,
        })

    return render(request, 'expertise/intervenants_list.html', {
        'medecin_rows': medecin_rows,
    })


@login_required(login_url='/login/')
def intervenant_history(request, pk):
    medecin = get_object_or_404(Medecin, pk=pk)
    history_map = _collect_medecin_histories(target_medecin=medecin)
    history = history_map.get(medecin.id, [])
    history.sort(key=lambda item: item['date'] or datetime.min, reverse=True)

    pending_entries = [entry for entry in history if entry.get('paiement') and not entry.get('invoice_number')]
    if request.GET.get('autogen') == '1' and pending_entries:
        _create_medecin_invoice(medecin, pending_entries)
        return redirect('intervenant_history', pk=medecin.pk)

    paid_entries = [entry for entry in history if entry.get('paiement')]
    total_brut = sum(entry.get('montant_brut', Decimal('0')) or Decimal('0') for entry in paid_entries)
    total_redevance = sum(entry.get('redevance', Decimal('0')) or Decimal('0') for entry in paid_entries)
    total_net = total_brut - total_redevance

    invoices = list(MedecinInvoice.objects.filter(medecin=medecin))
    latest_invoice = invoices[0] if invoices else None

    return render(request, 'expertise/intervenant_history.html', {
        'medecin': medecin,
        'history': history,
        'total_brut': total_brut,
        'total_redevance': total_redevance,
        'total_net': total_net,
        'latest_invoice': latest_invoice,
        'invoices': invoices,
        'pending_entries': pending_entries,
    })


@login_required(login_url='/login/')
def intervenant_invoice(request, pk, invoice_id):
    medecin = get_object_or_404(Medecin, pk=pk)
    invoice = get_object_or_404(MedecinInvoice, pk=invoice_id, medecin=medecin)

    pdf_file = _render_medecin_invoice_pdf(invoice)
    filename = f"FactureMedecin_{invoice.number}.pdf"

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@require_POST
@login_required(login_url='/login/')
def intervenant_invoice_delete(request, pk, invoice_id):
    medecin = get_object_or_404(Medecin, pk=pk)
    invoice = get_object_or_404(MedecinInvoice, pk=invoice_id, medecin=medecin)
    invoice.delete()
    return redirect('intervenant_history', pk=medecin.pk)


def _render_medecin_invoice_pdf(invoice):
    entries = [
        {
            'date': ligne.date_acte,
            'acte': ligne.act_label,
            'patient': f"{ligne.patient_prenom} {ligne.patient_nom}".strip(),
            'montant_brut': ligne.montant_brut,
            'redevance': ligne.redevance,
            'montant_net': ligne.montant_net,
        }
        for ligne in invoice.lignes.all().order_by('date_acte', 'pk')
    ]

    html_string = render_to_string('expertise/facture_medecin_historique_pdf.html', {
        'medecin': invoice.medecin,
        'invoice_number': invoice.number,
        'emission_date': invoice.emission_date,
        'entries': entries,
        'total_brut': invoice.total_brut,
        'total_redevance': invoice.total_redevance,
        'total_net': invoice.total_net,
    })

    return HTML(string=html_string).write_pdf()


@login_required(login_url='/login/')
def intervenant_invoices_zip(request, pk):
    medecin = get_object_or_404(Medecin, pk=pk)
    invoices = list(MedecinInvoice.objects.filter(medecin=medecin))

    if not invoices:
        return redirect('intervenant_history', pk=pk)

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as archive:
        for invoice in invoices:
            pdf_bytes = _render_medecin_invoice_pdf(invoice)
            filename = f"FactureMedecin_{invoice.number}.pdf"
            archive.writestr(filename, pdf_bytes)

    buffer.seek(0)
    zip_filename = f"FacturesMedecin_{medecin.nom}_{medecin.prenom}.zip".replace(' ', '_')

    response = HttpResponse(buffer.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{zip_filename}"'
    return response


@login_required(login_url='/login/')
def relance_factures(request):
    cutoff_date = timezone.now().date() - timedelta(days=180)

    impayes_qs = (
        FicheEvenement.objects
        .select_related('personnel__compagnie', 'bordereau', 'personnel')
        .filter(paiement=False, date_evenement__isnull=False, date_evenement__lte=cutoff_date)
        .order_by('personnel__compagnie__nom', 'bordereau__no_bordereau', 'no_facture')
    )

    company_rows = defaultdict(list)
    company_totals = defaultdict(int)
    company_refs = {}

    for evenement in impayes_qs:
        compagnie = evenement.personnel.compagnie if evenement.personnel else None
        key = compagnie.id if compagnie else 'none'

        company_refs[key] = compagnie
        company_rows[key].append({
            'bordereau_no': evenement.bordereau.no_bordereau if evenement.bordereau else '',
            'bordereau': evenement.bordereau,
            'facture': evenement.no_facture or 'N/A',
            'date_evenement': evenement.date_evenement,
            'total': evenement.total or 0,
        })
        company_totals[key] += evenement.total or 0

    def sort_key(k):
        compagnie = company_refs.get(k)
        if compagnie:
            nom = compagnie.nom.lower()
            iata = (compagnie.iata or '').lower()
        else:
            nom = 'zzz_compagnie_non_renseignee'
            iata = ''
        return (nom, iata)

    grouped_companies = []
    for key in sorted(company_rows.keys(), key=sort_key):
        compagnie = company_refs.get(key)
        rows = company_rows[key]
        rows_sorted = sorted(rows, key=lambda r: (r['bordereau_no'], r['facture']))
        grouped_companies.append({
            'compagnie': compagnie,
            'rows': rows_sorted,
            'total': company_totals[key],
        })

    total_global = sum(company_totals.values())

    return render(request, 'expertise/relance_factures.html', {
        'generated_at': timezone.now(),
        'cutoff_date': cutoff_date,
        'grouped_companies': grouped_companies,
        'total_global': total_global,
    })


from django.shortcuts import get_object_or_404, redirect


from django.shortcuts import get_object_or_404, redirect
from .models import Bordereau

def supprimer_bordereau(request, id):
    bordereau = get_object_or_404(Bordereau, id=id)

    # Dissocier les événements liés
    for evenement in bordereau.evenements.all():
        evenement.bordereau = None
        evenement.save()

    # Supprimer le bordereau
    bordereau.delete()

    return redirect('liste_bordereaux')

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from .models import FactureMedecin, Bordereau, FicheEvenement

from docx import Document
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from .models import Bordereau, Medecin, FactureMedecin, FicheEvenement

def telecharger_facture_medecin(request, bordereau_no, medecin_id):
    bordereau = get_object_or_404(Bordereau, no_bordereau=bordereau_no)
    medecin = get_object_or_404(Medecin, id=medecin_id)

    evenements = FicheEvenement.objects.filter(
        bordereau=bordereau
    ).filter(
        models.Q(medecin_cempn=medecin) |
        models.Q(medecin_oph=medecin) |
        models.Q(medecin_orl=medecin) |
        models.Q(medecin_radio=medecin) |
        models.Q(medecin_labo=medecin)
    )

    doc = Document()
    para = doc.add_heading('Centre Médical du Personnel Navigant de Polynésie française', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Dr. Christian Hellec', level=1)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('BP 380697 - 98718 Punaauia - Tahiti', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('Polynésie Française', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('mel : cmpnpf@gmail.com | Tel : +689.87.77.05.18 | Tel : +689.87.71.50.90', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('--------------------', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading(f"Honoraires pour le Dr {medecin.nom} {medecin.prenom}", level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para = doc.add_heading('--------------------', level=2)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")
    doc.add_paragraph(f"Bordereau : {bordereau_no}")

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "Patient"
    hdr_cells[2].text = "Montant brut"
    hdr_cells[3].text = "Redevance"
    hdr_cells[4].text = "Montant net"

    total = Decimal("0.00")
    total_redevance = Decimal("0.00")

    for e in evenements:
        if e.medecin_cempn == medecin:
            montant = e.honoraire_cempn or 0
        elif e.medecin_oph == medecin:
            montant = e.honoraire_cs_oph or 0
        elif e.medecin_orl == medecin:
            montant = e.honoraire_cs_orl or 0
        elif e.medecin_radio == medecin:
            montant = e.honoraire_cs_radio or 0
        elif e.medecin_labo == medecin:
            montant = (e.honoraire_cs_labo or 0) + (e.honoraire_cs_lbx or 0) + (e.honoraire_cs_toxique or 0)
        else:
            continue

        # Patient
        nom_patient = f"{e.personnel.prenom} {e.personnel.nom}"

        # Calculs
        montant_decimal = Decimal(montant)
        rate = _get_redevance_rate(medecin)
        redevance = (montant_decimal * rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        net = (montant_decimal - redevance).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        total += montant_decimal
        total_redevance += redevance

        row = table.add_row().cells
        row[0].text = str(e.date_evenement)
        row[1].text = nom_patient
        row[2].text = f"{montant_decimal:.0f} XPF"
        row[3].text = f"{redevance:.0f} XPF"
        row[4].text = f"{net:.0f} XPF"

    doc.add_paragraph("")
    doc.add_paragraph(f"Total brut : {total:.0f} XPF")
    doc.add_paragraph(f"Total redevance : {total_redevance:.0f} XPF")
    doc.add_paragraph(f"Total net à payer : {total - total_redevance:.0f} XPF")
    if medecin.iban:
        doc.add_paragraph(f"IBAN / RIB : {medecin.iban}")

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"Facture_{medecin.nom}_{bordereau_no}.docx".replace(" ", "_")
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    doc.save(response)

    return response


def telecharger_facture_medecin(request, bordereau_no, medecin_id):
    bordereau = get_object_or_404(Bordereau, no_bordereau=bordereau_no)
    medecin = get_object_or_404(Medecin, id=medecin_id)

    evenements_originaux = FicheEvenement.objects.filter(
        bordereau=bordereau
    ).filter(
        models.Q(medecin_cempn=medecin) |
        models.Q(medecin_oph=medecin) |
        models.Q(medecin_orl=medecin) |
        models.Q(medecin_radio=medecin) |
        models.Q(medecin_labo=medecin)
    )

    evenements = []
    total_brut = Decimal("0.00")
    total_redevance = Decimal("0.00")

    for e in evenements_originaux:
        if e.medecin_cempn == medecin:
            montant = e.honoraire_cempn or 0
        elif e.medecin_oph == medecin:
            montant = e.honoraire_cs_oph or 0
        elif e.medecin_orl == medecin:
            montant = e.honoraire_cs_orl or 0
        elif e.medecin_radio == medecin:
            montant = e.honoraire_cs_radio or 0
        elif e.medecin_labo == medecin:
            montant = sum(filter(None, [e.honoraire_cs_labo, e.honoraire_cs_lbx, e.honoraire_cs_toxique]))
        else:
            continue

        montant_decimal = Decimal(montant)
        rate = _get_redevance_rate(medecin)
        redevance = (montant_decimal * rate).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
        net = (montant_decimal - redevance).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

        evenements.append({
            "date": e.date_evenement.strftime('%d/%m/%Y'),
            "patient": f"{e.personnel.prenom} {e.personnel.nom}",
            "montant": f"{montant_decimal:.0f}",
            "redevance": f"{redevance:.0f}",
            "net": f"{net:.0f}",
        })

        total_brut += montant_decimal
        total_redevance += redevance

    total_net = total_brut - total_redevance

    html_string = render_to_string("expertise/facture_medecin_pdf.html", {
        "medecin": medecin,
        "bordereau": bordereau,
        "evenements": evenements,
        "total_brut": f"{total_brut:.0f}",
        "total_redevance": f"{total_redevance:.0f}",
        "total_net": f"{total_net:.0f}",
    })

    pdf_file = HTML(string=html_string).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Facture_{medecin.nom}_{bordereau.no_bordereau}.pdf"'
    return response

class CustomLoginView(LoginView):
    template_name = 'login.html'

class CustomLogoutView(LogoutView):
    next_page = '/login/'